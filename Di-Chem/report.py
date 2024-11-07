import csv
import ctypes
import sys
import tkinter as tk
from tkinter import filedialog
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows
import pandas as pd
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle, SimpleDocTemplate, PageBreak, Spacer
import keyboard
import platform as pf
import subprocess
import win32api
import win32api
import win32print
from PyPDF2 import PdfReader
import webbrowser
import os
from kivy.clock import Clock
from kivy.core.window import Window
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.dropdown import DropDown
from kivy.uix.label import Label
from kivy.uix.popup import Popup
from kivy.uix.textinput import TextInput
from kivy.uix.spinner import Spinner
from kivy.uix.button import Button
from kivy.lang import Builder
from kivymd.app import MDApp
from kivymd.toast import toast
from docx.oxml import OxmlElement
from kivymd.uix.button import MDFlatButton
from kivymd.uix.toolbar import MDTopAppBar
from kivymd.icon_definitions import md_icons
from kivymd.uix.dialog import MDDialog
from docx.shared import Inches
from kivy.uix.filechooser import FileChooserIconView, platform
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from kivy.app import App
from kivymd.uix.pickers import MDDatePicker,MDTimePicker
from database import *
import subprocess
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

report=Report()
chemical=Chemical()
batch_meta=BatchMetaData()


class OpenReportBatch(BoxLayout):
    def __init__(self, **kwargs):
        super(OpenReportBatch, self).__init__(**kwargs)
        self.selected_batch_name = None
        self.fetch_and_display_batch_names_report()

    def fetch_and_display_batch_names_report(self):
        threading.Thread(target=self.fetch_batch_names_async).start()

    def fetch_batch_names_async(self):
        try:
            batch_names = self.get_batch_names_report()
            last_100_batch_names = batch_names[-100:]
            Clock.schedule_once(lambda dt: self.display_batch_names(last_100_batch_names), 0)
        except Exception as e:
            print(f"Error fetching batch names: {e}")

    def display_batch_names(self, batch_names):
        batch_layout = self.ids.total_batch_report
        batch_layout.clear_widgets()  # Clear existing buttons

        for name in batch_names:
            if name:
                button = Button(
                    text=name, size_hint_y=None, height=30,
                    background_color=(3, 3, 3, 1), color=(0, 0, 0, 3)
                )
                button.bind(on_release=self.on_batch_button_click_report)
                batch_layout.add_widget(button)

    def get_batch_names_report(self):
        batch_names = batch_meta.fetch_batch_name()
        return [name for name in batch_names if name]

    def on_batch_button_click_report(self, instance):
        button_text = instance.text
        id_index = button_text.find("ID:")

        if id_index != -1:
            id_start = id_index + len("ID:")
            batchname_index = button_text.find("BatchName:", id_start)
            if batchname_index != -1:
                id_value = button_text[id_start:batchname_index].strip()
            else:
                id_value = button_text[id_start:].strip()

            selected_batch_label = self.ids.selected_batch_report
            selected_batch_label.text = id_value
            self.selected_batch_name = id_value

    def on_open_button_press_report(self, instance):
        selected_batch_label = self.ids.selected_batch_report
        if selected_batch_label.text:
            threading.Thread(target=self.fetch_and_display_selected_batch_report).start()
        else:
            toast("Error: No batch selected.", duration=0.6)

    def fetch_and_display_selected_batch_report(self):
        try:
            selected_batch_info = batch_meta.fetch_selected_batch_info(self.selected_batch_name)
            Clock.schedule_once(lambda dt: self.display_selected_batch_report(selected_batch_info), 0)
        except Exception as e:
            print(f"Error fetching selected batch info: {e}")

    def display_selected_batch_report(self, selected_batch_info):
        print(selected_batch_info)

        app = App.get_running_app()
        root = app.root
        App.get_running_app().root.display_selected_batch_data_report(self.selected_batch_name)
        root.display_selected_batch_report(selected_batch_info)
        self.fetch_and_display_batch_names_report()
        self.dismiss_popup()

    def dismiss_popup(self):
        parent_popup = self.parent.parent.parent  # Adjust the number of ".parent" based on your hierarchy
        if parent_popup:
            parent_popup.dismiss()

    def on_search_input_change_report(self, text):
        if text == '':
            self.fetch_and_display_batch_names_report()
        else:
            threading.Thread(target=self.update_batch_layout_report, args=(text,)).start()

    def update_batch_layout_report(self, search_text):
        try:
            all_batch_names = self.get_batch_names_report()
            filtered_batch_names = [name for name in all_batch_names if search_text.lower() in name.lower()]
            Clock.schedule_once(lambda dt: self.display_batch_names(filtered_batch_names), 0)
        except Exception as e:
            print(f"Error updating batch layout: {e}")

    def show_error_popup_report(self, message):
        error_popup = Popup(title='Error', content=Label(text=message), size_hint=(None, None), size=(300, 100))
        error_popup.open()


class MyApp(BoxLayout):

    def __init__(self, **kwargs):
        super(MyApp, self).__init__(**kwargs)
        self.chemical_names = self.get_chemical_names()
        self.batch_names=self.get_batch_names()
        self.batch_text_input=self.ids.batch_name_report
        self.text_input = self.ids.re_chem_spin
        self.dropdown = DropDown()
        self.populate_dropdown()
        self.ids.batch_meta_data.opacity=0
    def display_selected_batch_report(self, selected_batch_info):
        try:
            self.ids.re_chem_spin.text = ''
            self.ids.batch_name_report.text=''
            self.dropdown.dismiss()
            self.ids.batch_meta_data.opacity = 1
            self.ids.batch_id.text = str(selected_batch_info[0])
            self.ids.batch_name.text = str(selected_batch_info[1])
            self.ids.fabric_weight.text = str(selected_batch_info[2])
            self.ids.mlr.text = str(selected_batch_info[3])
            self.ids.machin_no.text = str(selected_batch_info[4])
            self.ids.created_by.text = str(selected_batch_info[5])
            self.ids.created_date.text = str(selected_batch_info[6])
        except KeyError as e:
            print(f"Error: ID not found - {e}")

    def display_batch_report(self, selected_batch_info):
        try:
            self.dropdown.dismiss()
            self.ids.batch_meta_data.opacity = 1
            self.ids.batch_id.text = str(selected_batch_info[0])
            self.ids.batch_name.text = str(selected_batch_info[1])
            self.ids.fabric_weight.text = str(selected_batch_info[2])
            self.ids.mlr.text = str(selected_batch_info[3])
            self.ids.machin_no.text = str(selected_batch_info[4])
            self.ids.created_by.text = str(selected_batch_info[5])
            self.ids.created_date.text = str(selected_batch_info[6])
        except KeyError as e:
            print(f"Error: ID not found - {e}")

    def display_selected_batch_data_report(self,selected_batch_name):

        data = report.fetch_report_data_batch(selected_batch_name)
        report_batch_grid = self.ids.report_batch_grid
        report_batch_grid.cols = 10
        report_batch_grid.clear_widgets()
        headers = ['ID', 'RecordID', 'BatchID', 'Chemical No','Chemical','UserName','Disp Wt', 'Water DispWt','DispDate', 'DispTime', ]

        for header in headers:
            report_batch_grid.add_widget(
                TextInput(text=header, halign='center', size_hint=(None, None), size=(193, 35),font_size=17,font_name="Aileron-Bold.otf"))

        for row in data:
            batch_id = str(row[0])
            button = Button(text=batch_id, size_hint=(None, None), size=(193, 30))

            report_batch_grid.add_widget(button)

            for value in row[1:]:
                text_input = TextInput(text=str(value), halign='center', size_hint=(None, None), size=(193, 30),
                                       readonly=True, background_color=(3, 3, 3, 1), multiline=False)
                report_batch_grid.add_widget(text_input)
    def get_batch_names(self):
        # Fetch chemical names from the database by calling the method
        batch_names = self.get_batch_names_report()
        batch_names = batch_meta.fetch_batch_name_report()

        # Remove empty strings from the list
        batch_names = [name for name in batch_names if name]

        # Log the chemical names for debugging
        print("batch names:", batch_names)

        return batch_names

    def on_batch_textinput_focus(self, instance, focused):
        if focused:
            self.show_batch_dropdown()
        else:
            self.dropdown.dismiss()

    def close_dropdown(self):
        self.dropdown.dismiss()

    def get_batch_names_report(self):
        batch_names = batch_meta.fetch_batch_report()
        return [name for name in batch_names if name]

    def show_batch_dropdown(self):
        self.dropdown.clear_widgets()
        batch_names = self.get_batch_names_report()
        for name in batch_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_batch(btn.text))
            self.dropdown.add_widget(btn)
        self.dropdown.max_height = 190  # Set the max_height directly on the dropdown object
        self.dropdown.open(self.batch_text_input if self.batch_text_input.focus else None)

    def populate_batch_dropdown(self):
        self.dropdown.clear_widgets()
        batch_names = self.get_batch_names_report()
        for name in batch_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_batch(btn.text))
            self.dropdown.add_widget(btn)

    def filter_batch_names(self, value):
        batch_names = self.get_batch_names_report()
        filtered_names = [name for name in batch_names if value.lower() in name.lower()]
        self.dropdown.clear_widgets()
        print("batchnames",batch_names)
        for name in batch_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_batch(btn.text))
            self.dropdown.add_widget(btn)
        if filtered_names:
            if not self.dropdown.parent:  # Check if the dropdown is already open
                self.dropdown.open(self.batch_text_input)
        else:
            self.dropdown.dismiss()

    def select_batch(self, name):
        self.batch_text_input.text = name
        self.dropdown.dismiss()


    def get_chemical_names(self):
        # Fetch chemical names from the database
        chemical_names = chemical.fetch_chemical_names()

        # Remove empty strings from the list
        chemical_names = [name for name in chemical_names if name]

        # Log the chemical names for debugging
        print("Chemical names:", chemical_names)

        return chemical_names

    def on_textinput_focus(self, instance, focused):
        if focused:
            self.show_dropdown()
        else:
            self.dropdown.dismiss()


    def show_dropdown(self):
        self.dropdown.clear_widgets()
        for name in self.chemical_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_chemical(btn.text))
            self.dropdown.add_widget(btn)
        self.dropdown.max_height = 190  # Set the max_height directly on the dropdown object
        self.dropdown.open(self.text_input if self.text_input.focus else None)

    def populate_dropdown(self):
        self.dropdown.clear_widgets()
        for name in self.chemical_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_chemical(btn.text))
            self.dropdown.add_widget(btn)

    def filter_chemicals(self, value):
        filtered_names = [name for name in self.chemical_names if value.lower() in name.lower()]
        self.dropdown.clear_widgets()
        for name in filtered_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_chemical(btn.text))
            self.dropdown.add_widget(btn)
        if filtered_names:
            if not self.dropdown.parent:  # Check if the dropdown is already open
                self.dropdown.open(self.text_input)
        else:
            self.dropdown.dismiss()

    def select_chemical(self, name):
        self.text_input.text = name
        self.dropdown.dismiss()

    def ReportData(self):
        self.ids.batch_meta_data.opacity = 0
        from_time = self.ids.selected_from_label.text
        to_time = self.ids.selected_to_label.text
        if not from_time.strip() or not to_time.strip() :
            toast("Error: No {} selected.".format(
                "from_time" if not from_time.strip() else
                "to_time"
            ), duration=0.6)
            return
        try:
            from_time_dt = datetime.strptime(from_time, '%Y-%m-%d %H:%M:%S')  # Adjust format if needed
            to_time_dt = datetime.strptime(to_time, '%Y-%m-%d %H:%M:%S')  # Adjust format if needed
        except ValueError:
            toast("Error: Invalid date format. Please use 'YYYY-MM-DD HH:MM:SS'.", duration=0.6)
            return

        # Get current time
        current_time = datetime.now()

        # Check if from_time or to_time is in the future
        if from_time_dt > current_time or to_time_dt > current_time:
            toast("Error: 'From Time' or 'To Time' is in the future.", duration=0.6)
            return
        data = report.fetch_report_data(from_time, to_time)
        report_batch_grid = self.ids.report_batch_grid
        report_batch_grid.cols=10
        report_batch_grid.clear_widgets()
        headers = ['ID', 'RecordID','BatchID','Chemical No','Chemical','UserName','Disp Wt','Water DispWt','DispDate', 'DispTime',]

        for header in headers:
            report_batch_grid.add_widget(TextInput(text=header, halign='center', size_hint=(None, None), size=(193, 35),font_size=17,font_name="Aileron-Bold.otf"))

        for row in data:
            batch_id = str(row[0])
            button = Button(text=batch_id, size_hint=(None, None), size=(193, 30))
            print(row)
            report_batch_grid.add_widget(button)

            for value in row[1:]:
                text_input = TextInput(text=str(value), halign='center', size_hint=(None, None), size=(193, 30),
                                            readonly=True, background_color=(3, 3, 3, 1), multiline=False)
                report_batch_grid.add_widget(text_input)


    def Chemical_data(self):
        self.ids.batch_meta_data.opacity = 0
        self.ids.batch_name_report.text=''
        self.dropdown.dismiss()
        from_time = self.ids.selected_from_label.text
        to_time = self.ids.selected_to_label.text
        chemical=self.ids.re_chem_spin.text
        if not from_time.strip() or not to_time.strip() or not chemical.strip():
            toast("Error: No {} selected.".format(
                "from_time" if not from_time.strip() else
                "to_time" if not to_time.strip() else
                "chemical"
            ), duration=0.6)
            return
        try:
            from_time_dt = datetime.strptime(from_time, '%Y-%m-%d %H:%M:%S')  # Adjust format if needed
            to_time_dt = datetime.strptime(to_time, '%Y-%m-%d %H:%M:%S')  # Adjust format if needed
        except ValueError:
            toast("Error: Invalid date format. Please use 'YYYY-MM-DD HH:MM:SS'.", duration=0.6)
            return

        # Get current time
        current_time = datetime.now()

        # Check if from_time or to_time is in the future
        if from_time_dt > current_time or to_time_dt > current_time:
            toast("Error: 'From Time' or 'To Time' is in the future.", duration=0.6)
            return
        data = report.fetch_chemical_data(from_time, to_time,chemical)
        report_batch_grid = self.ids.report_batch_grid
        report_batch_grid.clear_widgets()
        report_batch_grid.cols=10
        headers = ['ID', 'RecordID', 'BatchID', 'Chemical No','Chemical','UserName','Disp Wt','Water DispWt', 'DispDate', 'DispTime', ]

        for header in headers:
            report_batch_grid.add_widget(
                TextInput(text=header, halign='center', size_hint=(None, None), size=(193, 35),font_size=17,font_name="Aileron-Bold.otf"))



        for row in data:
            batch_id = str(row[0])
            button = Button(text=batch_id, size_hint=(None, None), size=(193, 30))
            print(row)
            report_batch_grid.add_widget(button)

            for value in row[1:]:
                text_input = TextInput(text=str(value), halign='center', size_hint=(None, None), size=(193, 30),
                                       readonly=True, background_color=(3, 3, 3, 1), multiline=False)
                report_batch_grid.add_widget(text_input)

    def bach_data(self):
        self.ids.batch_meta_data.opacity = 0
        self.ids.re_chem_spin.text = ''
        self.dropdown.dismiss()

        # Get text values
        from_time = self.ids.selected_from_label.text
        to_time = self.ids.selected_to_label.text
        batch_name = self.ids.batch_name_report.text
        print("batch_text:", batch_name)

        # Extract batch ID
        match = re.search(r'ID:\s*(\d+)', batch_name, re.IGNORECASE)
        batch_id = match.group(1) if match else None
        print(batch_id)

        # Check if any required field is empty
        if not from_time.strip() or not to_time.strip() or not batch_name.strip():
            toast("Error: No {} selected.".format(
                "from_time" if not from_time.strip() else
                "to_time" if not to_time.strip() else
                "batch_name"
            ), duration=0.6)
            return

        # Parse the from_time and to_time
        try:
            from_time_dt = datetime.strptime(from_time, '%Y-%m-%d %H:%M:%S')  # Adjust format if needed
            to_time_dt = datetime.strptime(to_time, '%Y-%m-%d %H:%M:%S')  # Adjust format if needed
        except ValueError:
            toast("Error: Invalid date format. Please use 'YYYY-MM-DD HH:MM:SS'.", duration=0.6)
            return

        # Get current time
        current_time = datetime.now()

        # Check if from_time or to_time is in the future
        if from_time_dt > current_time or to_time_dt > current_time:
            toast("Error: 'From Time' or 'To Time' is in the future.", duration=0.6)
            return

        # Fetch batch data
        data = report.fetch_batch_data_report(from_time, to_time, batch_id)
        selected_batch_info = batch_meta.fetch_selected_batch_info(batch_id)
        self.display_batch_report(selected_batch_info)

        # Update the grid
        report_batch_grid = self.ids.report_batch_grid
        report_batch_grid.clear_widgets()
        report_batch_grid.cols = 11
        headers = ['ID', 'RecordID', 'BatchID', 'Chemical No', 'Chemical','BatchName', 'UserName','Disp Wt', 'Water DispWt',
                   'DispDate', 'DispTime']

        for header in headers:
            report_batch_grid.add_widget(
                TextInput(text=header, halign='center', size_hint=(None, None), size=(175, 35), font_size=17,
                          font_name="Aileron-Bold.otf")
            )

        for row in data:
            batch_id = str(row[0])
            button = Button(text=batch_id, size_hint=(None, None), size=(175, 30))
            report_batch_grid.add_widget(button)

            for value in row[1:]:
                text_input = TextInput(
                    text=str(value), halign='center', size_hint=(None, None), size=(175, 30),
                    readonly=True, background_color=(3, 3, 3, 1), multiline=False
                )
                report_batch_grid.add_widget(text_input)
    def show_error_pop(self, message):
        # Create a popup with an error message
        error_popup = Popup(title='Error', content=Label(text=message), size_hint=(None, None), size=(300, 100))
        error_popup.open()
        self.current_popup = error_popup
    def openreport(self):
        self.ids.selected_from_label.text = "                                  "
        self.ids.selected_to_label.text = "                                  "
        self.ids.re_chem_spin.text = ""
        self.ids.batch_name_report.text = ''
        self.dropdown.dismiss()
        opendata = OpenReportBatch()
        openbatchWindow = Popup(title='Open Batch', content=opendata, size_hint=(None, None), size=(500, 400),
                                padding=10,title_size=20, title_font = "Aileron-Bold.otf")
        openbatchWindow.open()
class FileChooserDialog(BoxLayout):
    pass

class OpenGamil(BoxLayout):

    def __init__(self, **kwargs):
        super(OpenGamil, self).__init__(**kwargs)
        self.dropdown = DropDown()
        self.gmail_name_text = self.ids.gmail_name

    def to_gmail_textinput_focus(self,instance,focused):
        if focused:
            self.show_to_mail_dropdown()
        else:
            self.dropdown.dismiss()
    def show_to_mail_dropdown(self):
        self.dropdown.clear_widgets()
        to_mail_names = self.get_to_mail_names_json()
        for name in to_mail_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_togmail_id(btn.text))
            self.dropdown.add_widget(btn)
        self.dropdown.max_height = 190
        self.dropdown.open(self.ids.mail_to)
    def select_togmail_id(self,name):
        self.ids.mail_to.text = name
        self.dropdown.dismiss()

    def get_to_mail_names_json(self):
        to_mails=self.get_tomail_ids()
        total_to_mail=[name for name in to_mails if name]
        return total_to_mail

    def get_tomail_ids(self):
        file_path = 'email_id_data.json'
        if os.path.exists(file_path):
            with open(file_path, 'r') as file:
                data = json.load(file)
                # Extract 'To-Mail' if it exists in the entry
                emails = [entry.get("To-Mail") for entry in data if "To-Mail" in entry]
                return emails
        else:
            print("File not found.")
            return []

    def filter_gmail_to(self,value):
        email_name = self.get_tomail_ids()
        filtered_names = [name for name in email_name if value.lower() in name.lower()]

        self.dropdown.clear_widgets()
        for name in filtered_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_togmail_id(btn.text))
            self.dropdown.add_widget(btn)

        if filtered_names:
            if not self.dropdown.parent:  # Check if the dropdown is already open
                self.dropdown.open(self.ids.mail_to)
        else:
            self.dropdown.dismiss()
    def on_gmail_textinput_focus(self, instance, focused):
        if focused:
            self.show_gmail_dropdown()
        else:
            self.dropdown.dismiss()
    def show_gmail_dropdown(self):
        self.dropdown.clear_widgets()
        mail_names = self.getmail_names_json()
        for name in mail_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_gmail_id(btn.text))
            self.dropdown.add_widget(btn)
        self.dropdown.max_height = 190
        self.dropdown.open(self.gmail_name_text)
    def select_gmail_id(self,name):
        self.ids.gmail_name.text= name
        self.ids.gmail_password.text=self.fetch_password(name)
        self.dropdown.dismiss()


    def getmail_names_json(self):
        email_name = self.fetch_gmail_id()
        email_id=[name for name in email_name if name]
        return email_id

    def fetch_gmail_id(self):
        file_path = 'email_id_data.json'
        if os.path.exists(file_path):
            with open(file_path, 'r') as file:
                data = json.load(file)
                # Extract 'gmail' if it exists in the entry
                emails = [entry.get("gmail") for entry in data if "gmail" in entry]
                return emails
        else:
            print("File not found.")
            return []

    def fetch_password(self, email):
        file_path = 'email_id_data.json'
        if os.path.exists(file_path):
            with open(file_path, 'r') as file:
                data = json.load(file)
                # Search for the email and return the respective password
                for entry in data:
                    if entry["gmail"] == email:
                        return entry["password"]
                print("Email not found.")
                return None
        else:
            print("File not found.")
            return None

    def filter_gmail(self, value):
        email_name = self.fetch_gmail_id()
        filtered_names = [name for name in email_name if value.lower() in name.lower()]

        self.dropdown.clear_widgets()
        for name in filtered_names:
            btn = Button(text=name, size_hint_y=None, height=40)
            btn.bind(on_release=lambda btn: self.select_gmail_id(btn.text))
            self.dropdown.add_widget(btn)

        if filtered_names:
            if not self.dropdown.parent:  # Check if the dropdown is already open
                self.dropdown.open(self.ids.gmail_name)
        else:
            self.dropdown.dismiss()

    def export_gmail_parameters(self):
        try:
            if self.ids.mail_to.text == '':
                toast('correct your To mail Id' ,duration=0.6)
                return
            subject=self.ids.mail_subject.text
            body=self.ids.mail_body.text
            to_email=self.ids.mail_to.text
            filename ='report.pdf'
            if self.ids.gmail_name.text != '':
                gmail = self.ids.gmail_name.text
                print('gmail',gmail)
                password = self.ids.gmail_password.text
                print('password',password)
                self.send_email(subject,body,to_email,filename,gmail,password)
            else:
                gmail = 'dichemspi123@gmail.com'
                password='kftn taia qkvv onxu'
                self.send_email(subject, body, to_email, filename, gmail, password)
        except smtplib.SMTPRecipientsRefused as e:
            toast(f"Invalid To address",duration=0.6)
        except smtplib.SMTPAuthenticationError as e:
            toast("Authentication failed", duration=0.6)
        except smtplib.SMTPException as e:
            toast(f"An SMTP error occurred: {e}", duration=0.6)
        except Exception as e:
            toast(f"An error occurred: {e}", duration=0.6)

    def send_email(self,subject, body, to_email, filename, gmail, password):
        from_email = gmail

        msg = MIMEMultipart()
        msg['From'] = from_email
        msg['To'] = to_email
        msg['Subject'] = subject

        msg.attach(MIMEText(body, 'plain'))

        with open(filename, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f"attachment; filename={filename}")
            msg.attach(part)

        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)

        text = msg.as_string()
        server.sendmail(from_email, to_email, text)
        server.quit()

        toast(f"Email sent to {to_email}",duration=0.6)
        self.store_data(gmail,password)
        self.store_to_mails(to_email)
    def store_to_mails(self,to_mail):
        file_path = 'email_id_data.json'
        new_data = {'To-Mail': to_mail}

        if os.path.exists(file_path):
            with open(file_path, 'r') as file:
                data = json.load(file)
        else:
            data = []
        if new_data not in data:
            data.append(new_data)
            with open(file_path, 'w') as file:
                json.dump(data, file, indent=4)
            print("Data added successfully.")
        else:
            print("Data already exists. No new data added.")
    def store_data(self,email, password):
        file_path='email_id_data.json'
        new_data = {"gmail": email, "password": password}

        if os.path.exists(file_path):
            with open(file_path, 'r') as file:
                data = json.load(file)
        else:
            data = []
        if new_data not in data:
            data.append(new_data)
            with open(file_path, 'w') as file:
                json.dump(data, file, indent=4)
            print("Data added successfully.")
        else:
            print("Data already exists. No new data added.")
class ReportApp(MDApp):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.home_hotkey_id = None

    def build(self):
        self.icon = 'logindi.png'
        Window.maximize()
        Builder.load_file('reportpage.kv')
        sys.setrecursionlimit(10 ** 9)
        return MyApp()

    def switch_to_main(self):
        Builder.unload_file('reportpage.kv')
        print("Unloaded report")
        from main import MainApp
        ReportApp.stop(self)
        MainApp().run()

    def delay_fn(self, fn_name):
        self.root.ids.selected_from_label.disabled = True
        self.root.ids.selected_to_label.disabled = True
        self.root.ids.open_button.disabled = True
        self.root.ids.chemical_button.disabled = True
        self.root.ids.batch_button.disabled = True
        self.root.ids.all_button.disabled = True
        self.root.ids.home.disabled = True
        Clock.schedule_once(lambda dt: getattr(self, fn_name)(), 0.6)

    def change_button_color(self, button):
        button.md_bg_color = (77/255, 199/255, 226/255, 1)
        Clock.schedule_once(lambda dt: self.reset_button_color(button), .1)

    def reset_button_color(self, button):

        button.md_bg_color = (0.3, 0.3, 0.3, 1)
    def has_other_text(self,layout):
        headers = ['ID', 'RecordID', 'BatchID', 'Chemical No', 'Chemical', 'DispWt', 'Water DispWt', 'DispDate',
                   'DispTime']

        # Get all children of the GridLayout
        children = layout.children

        # Reverse the children to match the order of addition
        children = list(reversed(children))

        for child in children:
            if isinstance(child, TextInput):
                # Check if the text is not one of the headers
                if child.text not in headers:
                    return True
        return False
    def go_to_home(self,dt):
        self.switch_to_main()

    def show_save_dialog(self):
        report_batch_grid = self.root.ids.report_batch_grid

        if self.has_other_text(report_batch_grid):

            root = tk.Tk()
            icon_path = os.path.join(os.path.dirname(__file__), "logindi.ico")
            root.iconbitmap(default=icon_path)
            root.withdraw()
            filetypes = [
                ("PDF files", "*.pdf"),
                ("Word files", "*.docx"),
                ("PNG files", "*.png"),
                ("JPEG files", "*.jpeg"),
                ("Excel files", "*.xlsx"),
                ("CSV files",'*.csv')
            ]

            file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=filetypes)

            if file_path:
                if file_path.lower().endswith(".pdf"):
                    if self.root.ids.batch_name.text !='':
                        self.export_gridlayout_to_pdf(report_batch_grid, file_path)
                    else:
                        self.export_gridlayout_to_pdf_chem(report_batch_grid, file_path)
                elif file_path.lower().endswith(".docx"):
                    if self.root.ids.batch_name.text != '':
                        self.export_gridlayout_to_word(report_batch_grid, file_path)
                    else:
                        self.export_gridlayout_to_word_chem(report_batch_grid, file_path)
                elif file_path.lower().endswith(".png"):
                    if self.root.ids.batch_name.text != '':
                        self.export_gridlayout_to_image(report_batch_grid, file_path)
                    else:
                        self.export_gridlayout_to_image_chem(report_batch_grid, file_path)
                elif file_path.lower().endswith(".jpeg") or file_path.lower().endswith(".jpg"):
                    if self.root.ids.batch_name.text != '':
                        self.export_gridlayout_to_image(report_batch_grid, file_path)
                    else:
                        self.export_gridlayout_to_image_chem(report_batch_grid, file_path)
                elif file_path.lower().endswith(".xlsx"):
                    if self.root.ids.batch_name.text != '':
                        self.export_gridlayout_to_excel(report_batch_grid, file_path)
                    else:
                        self.export_gridlayout_to_excel_chem(report_batch_grid, file_path)
                elif file_path.lower().endswith(".csv"):
                    self.export_gridlayout_to_csv(report_batch_grid, file_path)
        else:
            toast("Nothing Report Opened", duration=0.6)

    def show_print_dialog(self):
        report_batch_grid = self.root.ids.report_batch_grid

        if self.has_other_text(report_batch_grid):

            report_batch_grid = self.root.ids.report_batch_grid
            root = tk.Tk()
            icon_path = os.path.join(os.path.dirname(__file__), "logindi.ico")
            root.iconbitmap(default=icon_path)
            root.withdraw()
            filetypes = [
                ("PDF files", "*.pdf")

            ]

            file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=filetypes)

            if file_path:
                if file_path.lower().endswith(".pdf"):
                    if self.root.ids.batch_name.text != '':
                        self.export_gridlayout_to_pdf(report_batch_grid, file_path)
                    else:
                        self.export_gridlayout_to_pdf_chem(report_batch_grid, file_path)

                    self.open_print_dialog(file_path)
        else:
            toast("Nothing Report Opened", duration=0.6)
    def open_print_dialog(self, file_path):
        if not os.path.isfile(file_path):
            print("File does not exist.")
            return

        system_name = pf.system()

        try:
            if system_name == "Windows":
                # Use win32print and win32api to send the file directly to the default printer
                printer_name = win32print.GetDefaultPrinter()
                win32api.ShellExecute(
                    0,
                    "print",
                    file_path,
                    f'/d:"{printer_name}"',
                    ".",
                    0
                )
            elif system_name == "Darwin":  # macOS
                subprocess.run(["lp", file_path])
            elif system_name == "Linux":
                subprocess.run(["lp", file_path])
            else:
                print("Unsupported operating system.")
                return

            # Show a toast notification
            toast("File sent to the default printer.", duration=0.6)
        except Exception as e:
            print(f"Failed to send file to the printer: {str(e)}")
            toast(f"Failed to send file to the printer: {str(e)}", duration=0.6)

    def export_gridlayout_to_csv(self, gridlayout, csv_filename):
        # Metadata grid data
        metadata_grid = self.root.ids.metadata_grid
        metadata_data = []
        metadata_cols = metadata_grid.cols
        metadata_cells = list(metadata_grid.children)
        metadata_cells.reverse()

        # Convert metadata grid cells to horizontal layout
        for row_index in range(0, len(metadata_cells), metadata_cols):
            row_data = []
            for col_index in range(metadata_cols):
                cell = metadata_cells[row_index + col_index]
                if col_index == 0:  # Remove colon only from labels in the first column
                    cell_text = cell.text.replace(':', '') if hasattr(cell, 'text') else ''
                else:
                    cell_text = cell.text if hasattr(cell, 'text') else ''
                row_data.append(cell_text)
            metadata_data.append(row_data)

        transposed_metadata_data = list(map(list, zip(*metadata_data)))

        # Report grid data
        report_data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()

        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(cell.text if hasattr(cell, 'text') else '')
            report_data.append(row_data)

        # Write to CSV file
        with open(csv_filename, 'w', newline='') as csvfile:
            csvwriter = csv.writer(csvfile)

            # Write metadata table
            for row in transposed_metadata_data:
                csvwriter.writerow(row)

            # Add an empty row between metadata and report tables
            csvwriter.writerow([])

            # Write report table
            for row in report_data:
                csvwriter.writerow(row)

        toast("CSV saved Successfully", duration=0.6)
    def export_gridlayout_to_pdf(self, gridlayout, pdf_filename, font_size=10):
        # Metadata grid data
        metadata_grid = self.root.ids.metadata_grid
        metadata_data = []
        metadata_cols = metadata_grid.cols
        metadata_cells = list(metadata_grid.children)
        metadata_cells.reverse()

        # Convert metadata grid cells to horizontal layout
        for row_index in range(0, len(metadata_cells), metadata_cols):
            row_data = []
            for col_index in range(metadata_cols):
                cell = metadata_cells[row_index + col_index]
                if col_index == 0:  # Remove colon only from labels in the first column
                    cell_text = cell.text.replace(':', '') if hasattr(cell, 'text') else ''
                else:
                    cell_text = cell.text if hasattr(cell, 'text') else ''
                row_data.append(cell_text)
            metadata_data.append(row_data)

        transposed_metadata_data = list(map(list, zip(*metadata_data)))

        # Report grid data
        report_data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()

        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(cell.text if hasattr(cell, 'text') else '')
            report_data.append(row_data)

        # Create PDF document
        doc = SimpleDocTemplate(pdf_filename, pagesize=A4)
        story = []

        # Create and style metadata table
        metadata_table = Table(transposed_metadata_data)
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), font_size),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(metadata_table)

        # Add a spacer between the tables
        story.append(Spacer(1, 12))

        # Create and style report table
        report_table = Table(report_data)
        report_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), font_size),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(report_table)

        # Build the PDF
        doc.build(story)
        toast("PDF saved Successfully", duration=0.6)

    def export_gridlayout_to_image(self, gridlayout, image_filename_prefix, default_cell_width=100, cell_height=30,
                                   font_path='arial.ttf', bold_font_path='arialbd.ttf', font_size=12, spacing=20):
        # Metadata grid data
        metadata_grid = self.root.ids.metadata_grid
        metadata_data = []
        metadata_cols = metadata_grid.cols
        metadata_cells = list(metadata_grid.children)
        metadata_cells.reverse()

        for row_index in range(0, len(metadata_cells), metadata_cols):
            row_data = []
            for col_index in range(metadata_cols):
                cell = metadata_cells[row_index + col_index]
                if col_index == 0:  # Remove colon only from labels in the first column
                    cell_text = cell.text.replace(':', '') if hasattr(cell, 'text') else ''
                else:
                    cell_text = cell.text if hasattr(cell, 'text') else ''
                row_data.append(cell_text)
            metadata_data.append(row_data)

        transposed_metadata_data = list(map(list, zip(*metadata_data)))

        # Report grid data
        data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()

        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(cell.text if hasattr(cell, 'text') else '')
            data.append(row_data)

        # Calculate column widths based on content
        def calculate_column_widths(data, font):
            col_widths = [default_cell_width] * len(data[0])
            for col_index in range(len(data[0])):
                max_width = default_cell_width
                for row in data:
                    text = row[col_index]
                    if text:
                        bbox = font.getbbox(text)
                        text_width = bbox[2] - bbox[0]
                        if text_width > max_width:
                            max_width = text_width
                col_widths[col_index] = max_width + 10  # Add some padding
            return col_widths

        regular_font = ImageFont.truetype(font_path, font_size)
        bold_font = ImageFont.truetype(bold_font_path, font_size)

        metadata_col_widths = calculate_column_widths(transposed_metadata_data, bold_font)
        report_col_widths = calculate_column_widths(data, regular_font)

        page_width = max(sum(metadata_col_widths), sum(report_col_widths))
        page_size = (page_width, 800)

        num_cols = len(data[0]) if data else 0
        num_rows = len(data)
        rows_per_page = (page_size[1] - cell_height - len(
            transposed_metadata_data) * cell_height - spacing) // cell_height  # Subtract height of metadata grid
        total_pages = (num_rows + rows_per_page - 1) // rows_per_page

        for page in range(total_pages):
            start_row = page * rows_per_page
            end_row = min(start_row + rows_per_page, num_rows)
            page_data = data[start_row:end_row]

            image = Image.new('RGB', page_size, color='white')
            draw = ImageDraw.Draw(image)

            if page == 0:
                # Calculate total width of the metadata table
                metadata_table_width = sum(metadata_col_widths)
                metadata_table_start_x = (page_width - metadata_table_width) / 2

                # Draw metadata table on the top, centered
                for row in range(len(transposed_metadata_data)):
                    for col in range(len(transposed_metadata_data[row])):
                        x0 = metadata_table_start_x + sum(metadata_col_widths[:col])
                        y0 = row * cell_height
                        x1 = x0 + metadata_col_widths[col]
                        y1 = y0 + cell_height

                        draw.rectangle([x0, y0, x1, y1], outline='black')

                        text = transposed_metadata_data[row][col]
                        font = bold_font if row == 0 else regular_font

                        bbox = font.getbbox(text)
                        text_width = bbox[2] - bbox[0]
                        text_height = bbox[3] - bbox[1]

                        text_x = x0 + (metadata_col_widths[col] - text_width) / 2
                        text_y = y0 + (cell_height - text_height) / 2
                        draw.text((text_x, text_y), text, font=font, fill='black')

                report_y_start = len(transposed_metadata_data) * cell_height + spacing
            else:
                report_y_start = 0

            # Draw report table below the metadata table with spacing
            for row in range(len(page_data)):
                for col in range(num_cols):
                    x0 = sum(report_col_widths[:col])
                    y0 = report_y_start + (row * cell_height)
                    x1 = x0 + report_col_widths[col]
                    y1 = y0 + cell_height

                    draw.rectangle([x0, y0, x1, y1], outline='black')

                    text = page_data[row][col]
                    font = bold_font if row == 0 else regular_font

                    bbox = font.getbbox(text)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]

                    text_x = x0 + (report_col_widths[col] - text_width) / 2
                    text_y = y0 + (cell_height - text_height) / 2
                    draw.text((text_x, text_y), text, font=font, fill='black')

            image_filename = f"{image_filename_prefix}_page_{page + 1}.png"
            image.save(image_filename)
            toast("Image saved Successfully", duration=0.6)

    def export_gridlayout_to_word(self, gridlayout, word_filename, cell_width=100, cell_height=30,
                                  font_path='arial.ttf', font_size=12, increase_factor=1.5):
        # Extract metadata from GridLayout
        metadata_grid = self.root.ids.metadata_grid
        metadata_data = []
        metadata_cols = metadata_grid.cols
        metadata_cells = list(metadata_grid.children)
        metadata_cells.reverse()

        for row_index in range(0, len(metadata_cells), metadata_cols):
            row_data = []
            for col_index in range(metadata_cols):
                cell = metadata_cells[row_index + col_index]
                if col_index == 0:  # Remove colon only from labels in the first column
                    cell_text = cell.text.replace(':', '') if hasattr(cell, 'text') else ''
                else:
                    cell_text = cell.text if hasattr(cell, 'text') else ''
                row_data.append(cell_text)
            metadata_data.append(row_data)

        transposed_metadata_data = list(map(list, zip(*metadata_data)))

        # Extract data from GridLayout
        data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()

        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(cell.text if hasattr(cell, 'text') else '')
            data.append(row_data)

        if cols == 9:
            page_width_inches = 904 / 96
            page_height_inches = 800 / 96
        elif cols == 10:
            page_width_inches = 1004 / 96
            page_height_inches = 800 / 96
        else:
            page_width_inches = 1004 / 96
            page_height_inches = 800 / 96

        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(page_width_inches)
        section.page_height = Inches(page_height_inches)

        metadata_col_widths = self.calculate_column_widths(transposed_metadata_data, cell_width=cell_width,
                                                           font_path=font_path, font_size=font_size)
        report_col_widths = self.calculate_column_widths(data, cell_width=cell_width, font_path=font_path,
                                                         font_size=font_size)

        num_metadata_cols = len(transposed_metadata_data[0]) if transposed_metadata_data else 0
        num_report_cols = len(data[0]) if data else 0

        # Center the metadata table on the first page
        if num_metadata_cols > 0:
            metadata_table = doc.add_table(rows=0, cols=num_metadata_cols)
            for row_index, row_data in enumerate(transposed_metadata_data):
                row = metadata_table.add_row().cells
                for col_index, cell_text in enumerate(row_data):
                    cell = row[col_index]
                    cell.text = cell_text
                    cell.width = Inches(metadata_col_widths[col_index] * increase_factor)

                    if row_index == 0:
                        cell.paragraphs[0].runs[0].font.bold = True

            self.set_table_borders(metadata_table)

        doc.add_paragraph()  # Add space between metadata and report tables

        # Create the report table
        if num_report_cols > 0:
            report_table = doc.add_table(rows=0, cols=num_report_cols)
            for row_index, row_data in enumerate(data):
                row = report_table.add_row().cells
                for col_index, cell_text in enumerate(row_data):
                    cell = row[col_index]
                    cell.text = cell_text
                    cell.width = Inches(report_col_widths[col_index] * increase_factor)

                    if row_index == 0:
                        cell.paragraphs[0].runs[0].font.bold = True

            self.set_table_borders(report_table)

        # Save the document
        doc.save(word_filename)
        toast("Word document saved Successfully", duration=0.6)

    def convert_to_number(self, value):
        try:
            return int(value)
        except ValueError:
            try:
                return float(value)
            except ValueError:
                return value

    def export_gridlayout_to_excel(self, gridlayout, excel_filename):
        # Extract data from GridLayout
        data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()

        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(self.convert_to_number(cell.text) if hasattr(cell, 'text') else '')
            data.append(row_data)

        # Extract metadata
        metadata_grid = self.root.ids.metadata_grid
        metadata_data = []
        metadata_cols = metadata_grid.cols
        metadata_cells = list(metadata_grid.children)
        metadata_cells.reverse()

        # Convert metadata grid cells to horizontal layout and apply conversion
        for row_index in range(0, len(metadata_cells), metadata_cols):
            row_data = []
            for col_index in range(metadata_cols):
                cell = metadata_cells[row_index + col_index]
                if col_index == 0:  # Remove colon only from labels in the first column
                    cell_text = cell.text.replace(':', '') if hasattr(cell, 'text') else ''
                else:
                    cell_text = cell.text if hasattr(cell, 'text') else ''
                row_data.append(self.convert_to_number(cell_text))  # Apply conversion here
            metadata_data.append(row_data)

        # Transpose metadata to fit as header
        metadata_header = list(map(list, zip(*metadata_data)))

        # Create Excel file
        wb = Workbook()
        ws = wb.active

        border = Border(left=Side(border_style='thin'),
                        right=Side(border_style='thin'),
                        top=Side(border_style='thin'),
                        bottom=Side(border_style='thin'))

        bold_font = Font(bold=True)
        regular_font = Font(bold=False)

        # Write metadata to Excel
        metadata_row_index = 1
        for row in metadata_header:
            ws.append(row)

        # Add a blank row for separation
        metadata_row_index += len(metadata_header)
        ws.append([''] * len(metadata_header[0]))  # Add an empty row for separation

        # Write report data to Excel below metadata
        report_start_row = metadata_row_index + 1
        df = pd.DataFrame(data)

        # Write the data rows of the report data
        for row in dataframe_to_rows(df, index=False, header=False):
            ws.append(row)

        # Remove all borders from the table
        max_col = len(metadata_header[0]) + len(df.columns)
        max_row = metadata_row_index + len(data) + 1

        # Remove borders from all cells
        for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
            for cell in row:
                cell.border = Border()  # Remove borders

        # Optionally, apply borders to specific sections (e.g., metadata headers)
        for row in ws.iter_rows(min_row=1, max_row=1, max_col=len(metadata_header[0])):
            for cell in row:
                cell.border = border
                cell.font = bold_font

        # Optionally, apply borders to the report data headers (if needed)
        if len(df.columns) > 0:
            for row in ws.iter_rows(min_row=report_start_row, max_row=report_start_row, max_col=len(df.columns)):
                for cell in row:
                    cell.border = border
                    cell.font = bold_font

        # Adjust column widths based on content
        def get_max_length(ws, col_idx):
            max_length = 0
            column_letter = get_column_letter(col_idx)
            for row in ws[column_letter]:
                try:
                    if len(str(row.value)) > max_length:
                        max_length = len(row.value)
                except:
                    pass
            return max_length

        # Adjust width for metadata columns
        for col_idx in range(1, len(metadata_header[0]) + 1):
            max_length = get_max_length(ws, col_idx)
            ws.column_dimensions[get_column_letter(col_idx)].width = max_length + 2  # Add some padding

        # Adjust width for report data columns
        start_col_idx = len(metadata_header[0]) + 1
        for col_idx in range(start_col_idx, start_col_idx + len(df.columns)):
            max_length = get_max_length(ws, col_idx)
            ws.column_dimensions[get_column_letter(col_idx)].width = max_length + 2  # Add some padding

        # Save the document
        wb.save(excel_filename)
        toast("Excel file saved successfully", duration=0.6)

    def export_gridlayout_to_pdf_chem(self, gridlayout, pdf_filename,font_size=10):
        metadata=self.root.ids.metadata_grid
        data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()
        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(cell.text if hasattr(cell, 'text') else '')
            data.append(row_data)

        doc = SimpleDocTemplate(pdf_filename, pagesize=A4)
        width, height = A4
        story = []

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), font_size),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(table)
        doc.build(story)
        toast("PDF saved Successfully",duration=0.6)

    def export_gridlayout_to_image_chem(self, gridlayout, image_filename_prefix, cell_width=100, cell_height=30,
                                   font_path='arial.ttf', bold_font_path='arialbd.ttf', font_size=12):
        data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()

        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(cell.text if hasattr(cell, 'text') else '')
            data.append(row_data)

        if cols == 9:
            page_size = (904, 800)
        elif cols == 10:
            page_size = (1004, 800)
        else:
            page_size = (1004, 800)

        num_cols = len(data[0]) if data else 0
        num_rows = len(data)
        rows_per_page = page_size[1] // cell_height
        total_pages = (num_rows + rows_per_page - 1) // rows_per_page

        regular_font = ImageFont.truetype(font_path, font_size)
        bold_font = ImageFont.truetype(bold_font_path, font_size)

        for page in range(total_pages):
            start_row = page * rows_per_page
            end_row = min(start_row + rows_per_page, num_rows)
            page_data = data[start_row:end_row]

            image = Image.new('RGB', page_size, color='white')
            draw = ImageDraw.Draw(image)

            for row in range(len(page_data)):
                for col in range(num_cols):
                    x0 = col * cell_width
                    y0 = row * cell_height
                    x1 = x0 + cell_width
                    y1 = y0 + cell_height

                    draw.rectangle([x0, y0, x1, y1], outline='black')

                    text = page_data[row][col]

                    if row == 0:
                        font = bold_font
                    else:
                        font = regular_font

                    bbox = font.getbbox(text)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]

                    text_x = x0 + (cell_width - text_width) / 2
                    text_y = y0 + (cell_height - text_height) / 2
                    draw.text((text_x, text_y), text, font=font, fill='black')

            image_filename = f"{image_filename_prefix}_page_{page + 1}.png"
            image.save(image_filename)
            toast("Image saved Successfully", duration=0.6)

    def calculate_column_widths(self, data, cell_width=100, font_path='arial.ttf', font_size=12):
        font = ImageFont.truetype(font_path, font_size)
        column_widths = []

        if data:
            num_cols = len(data[0])
            for col_index in range(num_cols):
                max_width = 0
                for row in data:
                    text = row[col_index]
                    bbox = font.getbbox(text)
                    text_width = bbox[2] - bbox[0]
                    if text_width > max_width:
                        max_width = text_width

                column_widths.append(cell_width / 96.0)

        return column_widths

    def set_table_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def export_gridlayout_to_word_chem(self, gridlayout, word_filename, cell_width=100, cell_height=30,
                                  font_path='arial.ttf', font_size=12, increase_factor=1.5):
        # Extract data from GridLayout
        data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()

        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(cell.text if hasattr(cell, 'text') else '')
            data.append(row_data)

        if cols == 9:
            page_width_inches = 904 / 96
            page_height_inches = 800 / 96
        elif cols == 10:
            page_width_inches = 1004 / 96
            page_height_inches = 800 / 96
        else:
            page_width_inches = 1004 / 96
            page_height_inches = 800 / 96

        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(page_width_inches)
        section.page_height = Inches(page_height_inches)

        column_widths = self.calculate_column_widths(data, cell_width=cell_width, font_path=font_path,
                                                     font_size=font_size)


        num_cols = len(data[0]) if data else 0
        table = doc.add_table(rows=0, cols=num_cols)

        for row_index, row_data in enumerate(data):
            row = table.add_row().cells
            for col_index, cell_text in enumerate(row_data):
                cell = row[col_index]
                cell.text = cell_text

                if row_index == 0:
                    cell.paragraphs[0].runs[0].font.bold = True

        # Adjust table width
        adjusted_widths = [cell_width / 96.0 * increase_factor for _ in
                           column_widths]

        for col_index, width in enumerate(adjusted_widths):
            for cell in table.columns[col_index].cells:
                cell.width = Inches(width)

        self.set_table_borders(table)

        # Save the document
        doc.save(word_filename)
        toast("Word document saved Successfully", duration=0.6)



    def export_gridlayout_to_excel_chem(self, gridlayout, excel_filename):

        data = []
        cols = gridlayout.cols
        cells = list(gridlayout.children)
        cells.reverse()

        for row_index in range(0, len(cells), cols):
            row_data = []
            for col_index in range(cols):
                cell = cells[row_index + col_index]
                row_data.append(cell.text if hasattr(cell, 'text') else '')
            data.append(row_data)

        df = pd.DataFrame(data)

        for col in df.columns:
            df[col] = df[col].apply(self.convert_to_number)

        wb = Workbook()
        ws = wb.active

        border = Border(left=Side(border_style=None),
                        right=Side(border_style=None),
                        top=Side(border_style=None),
                        bottom=Side(border_style=None))

        for row in dataframe_to_rows(df, index=False, header=False):
            ws.append(row)

        for row in ws.iter_rows(min_row=1, max_col=len(df.columns), max_row=len(data)):
            for cell in row:
                cell.border = border

        for col_idx, col in enumerate(df.columns):
            max_length = max(df[col].astype(str).map(len).max(), len(str(col)))
            ws.column_dimensions[get_column_letter(col_idx + 1)].width = max_length + 2  # Add some padding

        wb.save(excel_filename)
        toast("Excel file saved Successfully",duration=0.6)

    def on_start(self):
        # Register the hotkey
        report_batch_grid = self.root.ids.report_batch_grid
        report_batch_grid.clear_widgets()
        headers = ['ID', 'RecordID', 'BatchID', 'Chemical No','Chemical', 'DispWt','Water DispWt','DispDate', 'DispTime', ]

        for header in headers:
            report_batch_grid.add_widget(
                TextInput(text=header, halign='center', size_hint=(None, None), size=(214, 35),font_size=17,font_name="Aileron-Bold.otf"))
        # self.home_hotkey_id = keyboard.add_hotkey("home", self.on_home_pressed)





    def show_from_picker(self):
        date_dialog = MDDatePicker()
        date_dialog.bind(on_save=self.get_from_date)
        date_dialog.open()

    def get_from_date(self, instance, value, date_range):
        self.selected_date = value
        time_dialog = MDTimePicker()
        time_dialog.bind(on_save=self.get_date_from)
        time_dialog.open()

    def show_to_picker(self):
        date_dialog = MDDatePicker()
        date_dialog.bind(on_save=self.get_to_date)
        date_dialog.open()
    def get_to_date(self,instance, value, date_range):
        self.selected_date = value
        time_dialog = MDTimePicker()
        time_dialog.bind(on_save=self.get_date_to)
        time_dialog.open()

    def get_date_to(self, instance,time):
        selected_date_label = self.root.ids.selected_to_label
        formatted_datetime = f"{self.selected_date.strftime('%Y-%m-%d')} {time.strftime('%H:%M:%S')}"
        selected_date_label.text = f"{formatted_datetime}"

    def get_date_from(self, instance, time):
        selected_date_label = self.root.ids.selected_from_label
        formatted_datetime = f"{self.selected_date.strftime('%Y-%m-%d')} {time.strftime('%H:%M:%S')}"
        selected_date_label.text = f"{formatted_datetime}"
    def on_home_pressed(self, event=None):
        Clock.schedule_once(self.go_to_home)

    def gmail_popup(self):
        report_batch_grid = self.root.ids.report_batch_grid
        if self.has_other_text(report_batch_grid):

            self.pdf_generate_gmail()
            opengmail = OpenGamil()
            opengmailwindow = Popup(title='Gmail', content=opengmail, size_hint=(None, None), size=(600, 450),
                                    padding=10, title_size=25, title_font="Aileron-Bold.otf")
            opengmailwindow.open()
        else:
            toast("Nothing Report Opened", duration=0.6)

    def pdf_generate_gmail(self):

        try:
            file_path='report.pdf'
            report_batch_grid = self.root.ids.report_batch_grid
            if self.root.ids.batch_name.text != '':
                self.export_gridlayout_to_pdf(report_batch_grid, file_path)
            else:
                self.export_gridlayout_to_pdf_chem(report_batch_grid, file_path)

        except PermissionError as e:
            toast(f"Permission error: {e}", duration=0.6)
        except Exception as e:
            toast(f"An error occurred: {e}", duration=0.6)
    def __del__(self):
        pass

if __name__ == '__main__':
    ReportApp().run()
